from flask import Flask, request, send_file, jsonify
from flask_cors import CORS
from docx import Document
from io import BytesIO
import pandas as pd
import re

app = Flask(__name__)
CORS(app)

# Define formatting rules
ACRONYMS = {'CNC', 'RPM', 'HP', 'LNS', 'OSP', 'ATC', 'BT30', 'CAT40', 'APC'}

MODEL_CODES = {
    'LU15', 'RA-3F', 'HRF-155-LP', 'SL-25B', 'VF-2SS', 'VF-1', 'VF-2', 'VF-4', 'L25', 'T300', 'A-T14IA',
    'MC-600V-DC', 'CT500F', 'TS15', 'MC-800V62', 'VF-3', 'ES-450HII', 'MAM72-3VS', 'Cat40',
    'RA-2G', 'EA8', 'VX20', 'LB300-M', 'SJIII-3219', 'CONTURA 7/10/6', 'L370MW',
    'LT10-MY', 'EXCELLENCE 2F', 'ROBOFIL 240', 'VT500', 'GR-658N', 'L20 VIII',
    'VF-4SS', 'VF-9B/40', 'VA10M'
}

LOWERCASE_EXCEPTIONS = {'and', 'with', 'for', 'of', 'the', 'in'}

def smart_title_case(text):
    def fix_word(word):
        clean = word.strip(",.;:-")
        if clean.upper() in ACRONYMS or clean.upper() in MODEL_CODES:
            return clean.upper()
        if clean.lower() in LOWERCASE_EXCEPTIONS:
            return clean.lower()
        return clean.capitalize()

    words = re.split(r'(\s+)', text)  # Preserves spacing
    return ''.join(fix_word(word) if word.strip() else word for word in words)

# ---------------------------
# /upload → Preview only
# ---------------------------
@app.route('/upload', methods=['POST'])
def upload_preview():
    try:
        uploaded_file = request.files['file']
        template = request.form.get('template', 'standard')

        excel_data = pd.read_excel(uploaded_file, sheet_name=None)
        preview_lines = []

        for sheet_name, df in excel_data.items():
            if df.empty:
                continue

            df.fillna('', inplace=True)

            df.columns = [str(col).strip() for col in df.columns]

            col_make = df.columns[1]
            col_model = df.columns[2]
            col_year = df.columns[3]
            col_sn = df.columns[4]
            col_specs = df.columns[5]

            df['Year'] = df[col_year].astype(str).str.strip()
            df['Make'] = df[col_make].fillna('').astype(str).str.strip()
            df['Model'] = df[col_model].fillna('').astype(str).str.strip()
            df['Specs'] = df[col_specs].fillna('').apply(smart_title_case)
            df['Specs'] = df['Specs'].str.replace(r"S[/\\-]?N[:#]?\s*\w+", "", flags=re.IGNORECASE, regex=True).str.strip(",; ")
            df['Serial'] = df[col_sn].astype(str).str.strip().str.upper()

            grouped = df.groupby(['Year', 'Make', 'Model', 'Specs'])['Serial'].apply(list).reset_index()

            for _, row in grouped.iterrows():
                year, make, model, specs, serials = row['Year'], row['Make'], row['Model'], row['Specs'], row['Serial']
                quantity = len(serials)
                quantity_str = f"({quantity}) " if quantity > 1 else ""
                formatted_model = smart_title_case(f"{make} {model}")

                if template == "date_after_sn":
                    serial_str = f"S/N: {', '.join(serials).upper()}"
                    if year:
                        serial_str += f", New in {year}"
                    line = f"{quantity_str}{formatted_model}, {specs}, {serial_str}"
                else:
                    serial_str = f"S/N: {', '.join(serials).upper()}"
                    line = f"{quantity_str}{year} {formatted_model}, {specs}, {serial_str}"

                preview_lines.append(line)

        return jsonify({ "preview": preview_lines })

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

# ---------------------------
# /download → Word Export
# ---------------------------
@app.route('/download', methods=['POST'])
def download_file():
    try:
        uploaded_file = request.files['file']
        template = request.form.get('template', 'standard')

        excel_data = pd.read_excel(uploaded_file, sheet_name=None)
        doc = Document()
        doc.add_heading("Formatted Inventory", level=1)

        for sheet_name, df in excel_data.items():
            if df.empty:
                continue

            doc.add_heading(sheet_name, level=2)
            df.columns = [str(col).strip() for col in df.columns]

            col_make = df.columns[1]
            col_model = df.columns[2]
            col_year = df.columns[3]
            col_sn = df.columns[4]
            col_specs = df.columns[5]

            df['Year'] = df[col_year].astype(str).str.strip()
            df['Make'] = df[col_make].fillna('').astype(str).str.strip()
            df['Model'] = df[col_model].fillna('').astype(str).str.strip()
            df['Specs'] = df[col_specs].fillna('').apply(smart_title_case)
            df['Specs'] = df['Specs'].str.replace(r"S[/\\-]?N[:#]?\s*\w+", "", flags=re.IGNORECASE, regex=True).str.strip(",; ")
            df['Serial'] = df[col_sn].astype(str).str.strip().str.upper()

            grouped = df.groupby(['Year', 'Make', 'Model', 'Specs'])['Serial'].apply(list).reset_index()

            for _, row in grouped.iterrows():
                year, make, model, specs, serials = row['Year'], row['Make'], row['Model'], row['Specs'], row['Serial']
                quantity = len(serials)
                quantity_str = f"({quantity}) " if quantity > 1 else ""
                formatted_model = smart_title_case(f"{make} {model}")

                if template == "date_after_sn":
                    serial_str = f"S/N: {', '.join(serials).upper()}"
                    if year:
                        serial_str += f", New in {year}"
                    line = f"{quantity_str}{formatted_model}, {specs}, {serial_str}"
                else:
                    serial_str = f"S/N: {', '.join(serials).upper()}"
                    line = f"{quantity_str}{year} {formatted_model}, {specs}, {serial_str}"

                doc.add_paragraph(line)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return send_file(
            buffer,
            as_attachment=True,
            download_name="formatted_inventory.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

# ---------------------------
# Run server
# ---------------------------
if __name__ == '__main__':
    app.run(debug=True)